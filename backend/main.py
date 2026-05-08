"""
AI简历优化 - 后端服务
FastAPI + PyPDF2 文本提取 + DeepSeek API 智能分析
"""

import json
import os
import re
import urllib.request
from io import BytesIO
from typing import List

from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from PyPDF2 import PdfReader
from docx import Document

app = FastAPI(title="AI简历优化", version="1.0.0")

# CORS - allow frontend calls
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# DeepSeek API config
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
DEEPSEEK_BASE_URL = os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com/v1")

if not DEEPSEEK_API_KEY:
    print("[警告] DEEPSEEK_API_KEY 未设置！请设置环境变量或在 .env 文件中配置")


# ---------- Models ----------

class AnalyzeRequest(BaseModel):
    resume_text: str


class DownloadRequest(BaseModel):
    text: str
    format: str = "txt"  # "txt" or "docx"


class AnalyzeResponse(BaseModel):
    score: int
    summary: str
    suggestions: List[str]
    optimized: str


# ---------- PDF Text Extraction ----------

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file."""
    try:
        reader = PdfReader(BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        text = "\n".join(text_parts)
        # Clean up excessive whitespace
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r" {2,}", " ", text)
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"PDF解析失败: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a Word (.docx) file."""
    try:
        doc = Document(BytesIO(file_bytes))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        text = "\n".join(text_parts)
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Word解析失败: {str(e)}")


# ---------- DeepSeek AI Analysis ----------

RESUME_ANALYSIS_PROMPT = """你是一位资深HR和简历优化专家。请分析以下简历，并严格按JSON格式返回结果。

要求：
1. score: 综合评分0-100
2. summary: 100字以内的整体评价
3. suggestions: 3-5条具体优化建议
4. optimized: 优化后的完整简历文本

⚠️ 严格约束：优化版只能基于原文已有信息进行措辞优化和结构重组。绝对禁止编造以下内容：
- 不存在的工作经历、公司名称、职位
- 不存在的项目名称、项目成果、具体数据
- 不存在的技能、证书、学历信息
- 如果原文信息不足，宁可保留原文也不可编造

返回格式（只返回JSON，不要其他内容）：
{{
  "score": 85,
  "summary": "整体评价...",
  "suggestions": ["建议1", "建议2", "建议3"],
  "optimized": "优化后的完整简历..."
}}

以下是需要分析的简历：
---
{resume_text}
---"""


def analyze_resume_with_ai(resume_text: str) -> dict:
    """Send resume to DeepSeek for analysis using direct HTTP call."""
    try:
        prompt = RESUME_ANALYSIS_PROMPT.format(resume_text=resume_text)

        payload = json.dumps({
            "model": "deepseek-chat",
            "messages": [
                {
                    "role": "system",
                    "content": "你是资深HR和简历优化专家。只返回JSON。严格基于原文优化，不得编造任何新经历、项目、技能或数据。",
                },
                {"role": "user", "content": prompt},
            ],
            "temperature": 0.3,
            "max_tokens": 4096,
        }).encode("utf-8")

        url = f"{DEEPSEEK_BASE_URL}/chat/completions"
        req = urllib.request.Request(
            url,
            data=payload,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            },
        )

        resp = urllib.request.urlopen(req, timeout=60)
        body = json.loads(resp.read().decode("utf-8"))
        content = body["choices"][0]["message"]["content"]

        if not content:
            raise ValueError("AI返回为空")

        # Strip whitespace and try multiple strategies to extract JSON
        content = content.strip()

        # Strategy 1: Remove markdown code blocks if present
        if content.startswith("```"):
            lines = content.split("\n")
            # Remove opening ```json or ```
            if lines[0].startswith("```"):
                lines = lines[1:]
            # Remove closing ```
            if lines and lines[-1].strip() == "```":
                lines = lines[:-1]
            content = "\n".join(lines).strip()

        # Strategy 2: Find the outermost JSON object
        first_brace = content.find("{")
        last_brace = content.rfind("}")
        if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
            content = content[first_brace:last_brace + 1]

        result = json.loads(content)
        return result

    except json.JSONDecodeError as e:
        raise HTTPException(status_code=500, detail=f"AI返回格式解析失败: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI分析失败: {str(e)}")


# ---------- API Endpoints ----------

@app.get("/api/health")
async def health_check():
    return {"status": "ok", "service": "AI简历优化"}


@app.post("/api/upload")
async def upload_resume(file: UploadFile = File(...)):
    """Upload a PDF or Word resume and extract its text."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="文件名不能为空")

    filename_lower = file.filename.lower()
    if not (filename_lower.endswith(".pdf") or filename_lower.endswith(".docx")):
        raise HTTPException(status_code=400, detail="仅支持 PDF 和 Word (.docx) 文件")

    contents = await file.read()

    if len(contents) == 0:
        raise HTTPException(status_code=400, detail="文件为空")

    if len(contents) > 10 * 1024 * 1024:  # 10MB limit
        raise HTTPException(status_code=400, detail="文件过大，请上传小于10MB的文件")

    if filename_lower.endswith(".pdf"):
        text = extract_text_from_pdf(contents)
    else:
        text = extract_text_from_docx(contents)

    if not text or len(text) < 20:
        raise HTTPException(
            status_code=400,
            detail=f"简历内容太少（仅{len(text)}字），请上传完整简历",
        )

    return {"text": text, "length": len(text), "filename": file.filename}


@app.post("/api/analyze", response_model=AnalyzeResponse)
async def analyze_resume(request: AnalyzeRequest):
    """Analyze resume with AI."""
    if not request.resume_text or len(request.resume_text) < 20:
        raise HTTPException(status_code=400, detail="简历内容太少，请提供完整简历")

    result = analyze_resume_with_ai(request.resume_text)

    return AnalyzeResponse(
        score=result.get("score", 70),
        summary=result.get("summary", "无法生成评价"),
        suggestions=result.get("suggestions", []),
        optimized=result.get("optimized", request.resume_text),
    )


@app.post("/api/download")
async def download_resume(request: DownloadRequest):
    """Download optimized resume as TXT or DOCX."""
    from urllib.parse import quote

    if request.format == "docx":
        doc = Document()
        for line in request.text.split("\n"):
            if line.strip():
                doc.add_paragraph(line)
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        filename = quote("优化简历.docx")
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"},
        )
    else:
        buf = BytesIO(request.text.encode("utf-8"))
        filename = quote("优化简历.txt")
        return StreamingResponse(
            buf,
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"},
        )


if __name__ == "__main__":
    import uvicorn

    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
